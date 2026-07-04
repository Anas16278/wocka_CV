from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def add_horizontal_line(paragraph):
    """Add a horizontal line under a paragraph"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_font(run, size, bold=False, color=None):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def build_cv(data: dict, output_path: str):
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Name ─────────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(data['name'])
    set_font(name_run, 18, bold=True)
    name_para.paragraph_format.space_after = Pt(2)

    # ── Contact line ──────────────────────────────────────────────────────────
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run(f"{data['phone']} | {data['email']} | {data['location']}")
    set_font(contact_run, 10)
    contact_para.paragraph_format.space_after = Pt(8)

    # ── Helper: section heading with underline ────────────────────────────────
    def add_section(title):
        p = doc.add_paragraph()
        r = p.add_run(title)
        set_font(r, 12, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        add_horizontal_line(p)

    def add_body(text, bold=False, size=10.5):
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_font(r, size, bold=bold)
        p.paragraph_format.space_after = Pt(3)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(text)
        set_font(r, 10.5)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)

    # ── Personal Statement ────────────────────────────────────────────────────
    add_section("Personal Statement")
    add_body(data['personal_statement'])

    # ── Work Experience ───────────────────────────────────────────────────────
    add_section("Work Experience")
    p = doc.add_paragraph()
    r1 = p.add_run(f"{data['job_title']} – {data['company']}, London")
    set_font(r1, 10.5, bold=True)
    r2 = p.add_run(f"    |    {data['duration']}")
    set_font(r2, 10.5, color=(80, 80, 80))
    p.paragraph_format.space_after = Pt(3)

    for bullet in data['bullets']:
        add_bullet(bullet)

    # ── Education ─────────────────────────────────────────────────────────────
    add_section("Education")
    add_body("GCSEs – 2024", bold=True)
    for gcse in ["Mathematics – Grade 5", "English Language – Grade 5", "English Literature – Grade 5"]:
        add_bullet(gcse)
    add_body("T-Level in Digital Production, Design and Development (Sixth Form – Ongoing)", bold=True)

    # ── Skills ────────────────────────────────────────────────────────────────
    add_section("Skills")
    for skill in data['skills']:
        add_bullet(skill)

    # ── References ────────────────────────────────────────────────────────────
    add_section("References")
    add_body("Available upon request.")

    doc.save(output_path)


def build_cover_letter(data: dict, output_path: str):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Name
    p = doc.add_paragraph()
    r = p.add_run(data['name'])
    r.font.name = 'Calibri'
    r.font.size = Pt(14)
    r.font.bold = True
    p.paragraph_format.space_after = Pt(2)

    # Contact
    p = doc.add_paragraph()
    r = p.add_run(f"{data['phone']} | {data['email']}")
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(20)

    # Cover letter paragraphs
    for para_text in data['cover_letter'].split('\n\n'):
        text = para_text.strip()
        if not text:
            continue
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(10)

    # Sign off
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Yours sincerely,")
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(20)

    p = doc.add_paragraph()
    r = p.add_run(data['name'])
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.font.bold = True

    doc.save(output_path)
